#!/usr/bin/env python3
"""
Test script for DOCX to PDF conversion
"""

import os
from pathlib import Path
from file_converter import FileConversionService

def create_test_docx():
    """Create a simple test DOCX with multiple paragraphs and formatting"""
    try:
        import docx
        from docx.shared import Inches
        
        # Create a new document
        doc = docx.Document()
        
        # Add title
        title = doc.add_heading('Test Document for PDF Conversion', 0)
        
        # Add some paragraphs with different styles
        doc.add_heading('Introduction', level=1)
        intro_para = doc.add_paragraph(
            'This is a test DOCX document created to verify the DOCX to PDF conversion functionality. '
            'It contains multiple paragraphs, headings, and formatted text to ensure proper conversion.'
        )
        
        doc.add_heading('Main Content', level=1)
        
        # Add multiple paragraphs
        paragraphs = [
            'This is the first paragraph of the main content. It demonstrates basic text conversion.',
            'This is the second paragraph with some additional content to test multi-paragraph conversion.',
            'The third paragraph includes more text to verify that the PDF conversion maintains proper formatting and spacing.',
        ]
        
        for para_text in paragraphs:
            doc.add_paragraph(para_text)
        
        # Add a subheading
        doc.add_heading('Technical Details', level=2)
        doc.add_paragraph(
            'The conversion process should maintain the document structure, including headings, '
            'paragraphs, and basic formatting. Each paragraph should appear as a separate element '
            'in the resulting PDF document.'
        )
        
        # Add a simple table
        doc.add_heading('Sample Table', level=2)
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        
        # Add table content
        cells = [
            ['Feature', 'Status'],
            ['DOCX Reading', 'Working'],
            ['PDF Generation', 'Testing']
        ]
        
        for i, row_data in enumerate(cells):
            row = table.rows[i]
            for j, cell_data in enumerate(row_data):
                row.cells[j].text = cell_data
        
        # Add conclusion
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph(
            'This test document should convert successfully to PDF format, maintaining the structure '
            'and readability of the original DOCX content.'
        )
        
        # Save the document
        test_docx_path = "test_document.docx"
        doc.save(test_docx_path)
        
        print(f"✅ Created test DOCX: {test_docx_path}")
        return test_docx_path
        
    except Exception as e:
        print(f"❌ Failed to create test DOCX: {e}")
        return None

def test_docx_to_pdf():
    """Test DOCX to PDF conversion"""
    print("🧪 Testing DOCX to PDF Conversion")
    print("=" * 40)
    
    # Create test DOCX
    test_docx = create_test_docx()
    if not test_docx:
        return False
    
    try:
        # Initialize conversion service
        service = FileConversionService()
        
        # Check supported formats
        formats = service.list_supported_formats()
        print("📋 Document converter formats:")
        doc_formats = formats.get('document', {})
        print(f"   Input: {doc_formats.get('input', [])}")
        print(f"   Output: {doc_formats.get('output', [])}")
        
        # Test conversion to PDF
        output_pdf = "test_output.pdf"
        
        print(f"📄 Converting {test_docx} to PDF...")
        success = service.convert_file(test_docx, output_pdf)
        
        if success:
            print(f"✅ Conversion successful!")
            print(f"📦 Output PDF file: {output_pdf}")
            
            # Check if PDF file exists and has content
            if os.path.exists(output_pdf):
                file_size = os.path.getsize(output_pdf)
                print(f"📊 PDF file size: {file_size} bytes")
                
                if file_size > 0:
                    print("✅ PDF file created with content")
                    
                    # Try to verify it's a valid PDF by reading the header
                    with open(output_pdf, 'rb') as f:
                        header = f.read(8)
                        if header.startswith(b'%PDF-'):
                            print("✅ Valid PDF header detected")
                            return True
                        else:
                            print("❌ Invalid PDF header")
                            return False
                else:
                    print("❌ PDF file is empty")
                    return False
            else:
                print("❌ Output PDF file not found")
                return False
        else:
            print("❌ Conversion failed")
            return False
            
    except Exception as e:
        print(f"❌ Test failed: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # Cleanup
        for file in [test_docx, "test_output.pdf"]:
            if os.path.exists(file):
                try:
                    os.remove(file)
                    print(f"🧹 Cleaned up: {file}")
                except:
                    pass

if __name__ == "__main__":
    success = test_docx_to_pdf()
    if success:
        print("\n🎉 DOCX to PDF conversion test PASSED!")
    else:
        print("\n❌ DOCX to PDF conversion test FAILED!")