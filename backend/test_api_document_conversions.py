#!/usr/bin/env python3
"""
Test script for document conversions via API
Tests both PDF to images and DOCX to PDF conversions
"""

import requests
import os
import tempfile
from pathlib import Path

def create_test_docx():
    """Create a simple test DOCX"""
    try:
        import docx
        
        doc = docx.Document()
        doc.add_heading('API Test Document', 0)
        doc.add_paragraph('This is a test DOCX document for API conversion testing.')
        doc.add_paragraph('It should convert successfully to PDF format.')
        
        test_docx_path = "test_api_document.docx"
        doc.save(test_docx_path)
        
        print(f"✅ Created test DOCX: {test_docx_path}")
        return test_docx_path
        
    except Exception as e:
        print(f"❌ Failed to create test DOCX: {e}")
        return None

def create_test_pdf():
    """Create a simple test PDF"""
    try:
        import fitz
        
        doc = fitz.open()
        for page_num in range(2):
            page = doc.new_page()
            text = f"API Test Page {page_num + 1}\n\nThis is a test PDF for API conversion."
            page.insert_text((50, 100), text, fontsize=16, color=(0, 0, 0))
        
        test_pdf_path = "test_api_document.pdf"
        doc.save(test_pdf_path)
        doc.close()
        
        print(f"✅ Created test PDF: {test_pdf_path}")
        return test_pdf_path
        
    except Exception as e:
        print(f"❌ Failed to create test PDF: {e}")
        return None

def test_api_health():
    """Test API health"""
    try:
        response = requests.get("http://localhost:5000/api/health", timeout=5)
        if response.status_code == 200:
            print("✅ API server is healthy")
            return True
        else:
            print(f"❌ API health check failed: {response.status_code}")
            return False
    except Exception as e:
        print(f"❌ API health check failed: {e}")
        return False

def test_docx_to_pdf_api():
    """Test DOCX to PDF conversion via API"""
    print("\n📄 Testing DOCX to PDF conversion via API...")
    
    test_docx = create_test_docx()
    if not test_docx:
        return False
    
    try:
        with open(test_docx, 'rb') as f:
            files = {'file': f}
            data = {
                'source_format': 'DOCX',
                'target_format': 'PDF'
            }
            
            response = requests.post(
                "http://localhost:5000/api/convert",
                files=files,
                data=data,
                timeout=30
            )
        
        if response.status_code == 200:
            result = response.json()
            if result['success']:
                print("✅ DOCX to PDF API conversion successful!")
                print(f"📦 Converted filename: {result['converted_filename']}")
                print(f"📊 File size: {result['size']} bytes")
                
                # Test download
                download_url = f"http://localhost:5000{result['download_url']}"
                download_response = requests.get(download_url, timeout=10)
                
                if download_response.status_code == 200:
                    print("✅ PDF download successful!")
                    
                    # Verify it's a valid PDF
                    if download_response.content.startswith(b'%PDF-'):
                        print("✅ Valid PDF file downloaded")
                        return True
                    else:
                        print("❌ Downloaded file is not a valid PDF")
                        return False
                else:
                    print(f"❌ PDF download failed: {download_response.status_code}")
                    return False
            else:
                print(f"❌ DOCX to PDF conversion failed: {result.get('error', 'Unknown error')}")
                return False
        else:
            print(f"❌ API request failed: {response.status_code}")
            return False
            
    except Exception as e:
        print(f"❌ DOCX to PDF test failed: {e}")
        return False
    finally:
        if os.path.exists(test_docx):
            os.remove(test_docx)

def test_pdf_to_images_api():
    """Test PDF to images conversion via API"""
    print("\n📄 Testing PDF to images conversion via API...")
    
    test_pdf = create_test_pdf()
    if not test_pdf:
        return False
    
    try:
        with open(test_pdf, 'rb') as f:
            files = {'file': f}
            data = {
                'source_format': 'PDF',
                'target_format': 'JPEG'
            }
            
            response = requests.post(
                "http://localhost:5000/api/convert",
                files=files,
                data=data,
                timeout=30
            )
        
        if response.status_code == 200:
            result = response.json()
            if result['success']:
                print("✅ PDF to images API conversion successful!")
                print(f"📦 Converted filename: {result['converted_filename']}")
                print(f"📊 File size: {result['size']} bytes")
                
                # Test download
                download_url = f"http://localhost:5000{result['download_url']}"
                download_response = requests.get(download_url, timeout=10)
                
                if download_response.status_code == 200:
                    print("✅ ZIP download successful!")
                    
                    # Save and verify it's a valid ZIP
                    with open("test_downloaded.zip", "wb") as f:
                        f.write(download_response.content)
                    
                    import zipfile
                    try:
                        with zipfile.ZipFile("test_downloaded.zip", 'r') as zf:
                            files = zf.namelist()
                            print(f"📁 ZIP contains {len(files)} files:")
                            for file in files:
                                print(f"   - {file}")
                        
                        os.remove("test_downloaded.zip")
                        return True
                    except zipfile.BadZipFile:
                        print("❌ Downloaded file is not a valid ZIP")
                        return False
                else:
                    print(f"❌ ZIP download failed: {download_response.status_code}")
                    return False
            else:
                print(f"❌ PDF to images conversion failed: {result.get('error', 'Unknown error')}")
                return False
        else:
            print(f"❌ API request failed: {response.status_code}")
            return False
            
    except Exception as e:
        print(f"❌ PDF to images test failed: {e}")
        return False
    finally:
        if os.path.exists(test_pdf):
            os.remove(test_pdf)

def main():
    """Main test function"""
    print("🧪 Testing Document Conversions via API")
    print("=" * 50)
    
    # Test API health first
    if not test_api_health():
        print("\n❌ API server is not available!")
        print("Please start the API server: python api_server.py")
        return False
    
    # Test both conversions
    docx_success = test_docx_to_pdf_api()
    pdf_success = test_pdf_to_images_api()
    
    print("\n" + "=" * 50)
    print("📊 Test Results:")
    print(f"   DOCX to PDF: {'✅ PASSED' if docx_success else '❌ FAILED'}")
    print(f"   PDF to Images: {'✅ PASSED' if pdf_success else '❌ FAILED'}")
    
    if docx_success and pdf_success:
        print("\n🎉 All document conversion tests PASSED!")
        return True
    else:
        print("\n❌ Some document conversion tests FAILED!")
        return False

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)