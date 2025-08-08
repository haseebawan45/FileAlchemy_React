#!/usr/bin/env python3
"""
Comprehensive test script for text conversions
Tests TXT, HTML, DOCX, and PDF conversions
"""

import os
from pathlib import Path
from file_converter import FileConversionService

def create_test_txt():
    """Create a test text file"""
    content = """SAMPLE TEXT DOCUMENT

Introduction
This is a sample text document for testing conversions.
It contains multiple paragraphs and different types of content.

MAIN CONTENT
Here is the main content of the document.
This paragraph demonstrates normal text formatting.

Another paragraph with more content to test the conversion process.
The text should maintain its structure and readability.

Technical Details
- This document tests text conversion functionality
- It includes headings, paragraphs, and lists
- The conversion should preserve the document structure

Conclusion
This test document should convert successfully to various formats
while maintaining readability and structure."""
    
    test_file = "test_sample.txt"
    with open(test_file, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"✅ Created test TXT: {test_file}")
    return test_file

def create_test_html():
    """Create a test HTML file"""
    content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Sample HTML Document</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 20px; }
        h1 { color: #333; }
        h2 { color: #666; }
        p { line-height: 1.6; }
    </style>
</head>
<body>
    <h1>Sample HTML Document</h1>
    
    <h2>Introduction</h2>
    <p>This is a sample HTML document for testing conversions.</p>
    <p>It contains various HTML elements including headings, paragraphs, and formatting.</p>
    
    <h2>Main Content</h2>
    <p>Here is the main content of the document with <strong>bold text</strong> and <em>italic text</em>.</p>
    <p>This paragraph demonstrates HTML formatting that should be converted to plain text.</p>
    
    <h2>Lists and Structure</h2>
    <ul>
        <li>First list item</li>
        <li>Second list item</li>
        <li>Third list item</li>
    </ul>
    
    <h2>Conclusion</h2>
    <p>This test document should convert successfully to text and other formats.</p>
</body>
</html>"""
    
    test_file = "test_sample.html"
    with open(test_file, 'w', encoding='utf-8') as f:
        f.write(content)
    
    print(f"✅ Created test HTML: {test_file}")
    return test_file

def create_test_docx():
    """Create a test DOCX file"""
    try:
        import docx
        
        doc = docx.Document()
        doc.add_heading('Sample DOCX Document', 0)
        
        doc.add_heading('Introduction', level=1)
        doc.add_paragraph('This is a sample DOCX document for testing conversions.')
        doc.add_paragraph('It contains headings, paragraphs, and formatted content.')
        
        doc.add_heading('Main Content', level=1)
        doc.add_paragraph('Here is the main content with various formatting options.')
        doc.add_paragraph('This paragraph tests the conversion from DOCX to other formats.')
        
        doc.add_heading('Conclusion', level=1)
        doc.add_paragraph('This document should convert successfully to text and PDF formats.')
        
        test_file = "test_sample.docx"
        doc.save(test_file)
        
        print(f"✅ Created test DOCX: {test_file}")
        return test_file
        
    except Exception as e:
        print(f"❌ Failed to create test DOCX: {e}")
        return None

def test_conversion(service, input_file, output_file, conversion_name):
    """Test a single conversion"""
    print(f"  📄 Testing {conversion_name}...")
    
    if not os.path.exists(input_file):
        print(f"  ❌ Input file not found: {input_file}")
        return False
    
    success = service.convert_file(input_file, output_file)
    
    if success and os.path.exists(output_file):
        file_size = os.path.getsize(output_file)
        print(f"  ✅ {conversion_name} successful! Output: {file_size} bytes")
        return True
    else:
        print(f"  ❌ {conversion_name} failed!")
        return False

def test_all_text_conversions():
    """Test all text conversions"""
    print("🧪 Testing All Text Conversions")
    print("=" * 50)
    
    # Initialize conversion service
    service = FileConversionService()
    
    # Check supported formats
    formats = service.list_supported_formats()
    print("📋 Document converter formats:")
    doc_formats = formats.get('document', {})
    print(f"   Input: {doc_formats.get('input', [])}")
    print(f"   Output: {doc_formats.get('output', [])}")
    print()
    
    # Create test files
    test_txt = create_test_txt()
    test_html = create_test_html()
    test_docx = create_test_docx()
    
    results = {}
    test_files = []
    
    try:
        # Test TXT conversions
        print("📝 Testing TXT Conversions:")
        if test_txt:
            results['TXT→HTML'] = test_conversion(service, test_txt, "output_txt_to_html.html", "TXT to HTML")
            test_files.append("output_txt_to_html.html")
            
            results['TXT→PDF'] = test_conversion(service, test_txt, "output_txt_to_pdf.pdf", "TXT to PDF")
            test_files.append("output_txt_to_pdf.pdf")
            
            if test_docx:  # Only test if DOCX creation worked
                results['TXT→DOCX'] = test_conversion(service, test_txt, "output_txt_to_docx.docx", "TXT to DOCX")
                test_files.append("output_txt_to_docx.docx")
        
        print()
        
        # Test HTML conversions
        print("🌐 Testing HTML Conversions:")
        if test_html:
            results['HTML→TXT'] = test_conversion(service, test_html, "output_html_to_txt.txt", "HTML to TXT")
            test_files.append("output_html_to_txt.txt")
            
            results['HTML→PDF'] = test_conversion(service, test_html, "output_html_to_pdf.pdf", "HTML to PDF")
            test_files.append("output_html_to_pdf.pdf")
        
        print()
        
        # Test DOCX conversions
        print("📄 Testing DOCX Conversions:")
        if test_docx:
            results['DOCX→TXT'] = test_conversion(service, test_docx, "output_docx_to_txt.txt", "DOCX to TXT")
            test_files.append("output_docx_to_txt.txt")
            
            results['DOCX→PDF'] = test_conversion(service, test_docx, "output_docx_to_pdf.pdf", "DOCX to PDF")
            test_files.append("output_docx_to_pdf.pdf")
        
        print()
        
        # Test results summary
        print("📊 Conversion Results:")
        print("-" * 30)
        
        total_tests = len(results)
        successful_tests = sum(1 for success in results.values() if success)
        
        for conversion, success in results.items():
            status = "✅ PASSED" if success else "❌ FAILED"
            print(f"   {conversion}: {status}")
        
        print("-" * 30)
        print(f"   Total: {successful_tests}/{total_tests} conversions successful")
        
        # Verify some output files
        print("\n🔍 Verifying Output Files:")
        
        # Check PDF files
        pdf_files = [f for f in test_files if f.endswith('.pdf')]
        for pdf_file in pdf_files:
            if os.path.exists(pdf_file):
                with open(pdf_file, 'rb') as f:
                    header = f.read(8)
                    if header.startswith(b'%PDF-'):
                        print(f"   ✅ {pdf_file}: Valid PDF")
                    else:
                        print(f"   ❌ {pdf_file}: Invalid PDF")
        
        # Check HTML files
        html_files = [f for f in test_files if f.endswith('.html')]
        for html_file in html_files:
            if os.path.exists(html_file):
                with open(html_file, 'r', encoding='utf-8') as f:
                    content = f.read()
                    if '<!DOCTYPE html>' in content or '<html' in content:
                        print(f"   ✅ {html_file}: Valid HTML")
                    else:
                        print(f"   ❌ {html_file}: Invalid HTML")
        
        return successful_tests == total_tests
        
    except Exception as e:
        print(f"❌ Test failed with error: {e}")
        import traceback
        traceback.print_exc()
        return False
        
    finally:
        # Cleanup
        cleanup_files = [test_txt, test_html, test_docx] + test_files
        for file in cleanup_files:
            if file and os.path.exists(file):
                try:
                    os.remove(file)
                    print(f"🧹 Cleaned up: {file}")
                except:
                    pass

if __name__ == "__main__":
    success = test_all_text_conversions()
    if success:
        print("\n🎉 All text conversion tests PASSED!")
    else:
        print("\n❌ Some text conversion tests FAILED!")
    
    exit(0 if success else 1)