#!/usr/bin/env python3
"""
Final integration test for web interface
Tests key conversions through the API to ensure web compatibility
"""

import requests
import os
import tempfile
from pathlib import Path

def create_test_files():
    """Create test files for web integration testing"""
    files = {}
    
    # Create test TXT
    with open("web_test.txt", 'w', encoding='utf-8') as f:
        f.write("Web Integration Test\n\nThis is a test document for web interface validation.\nIt should convert successfully through the API.")
    files['txt'] = "web_test.txt"
    
    # Create test HTML
    html_content = """<!DOCTYPE html>
<html><head><title>Web Test</title></head>
<body>
<h1>Web Integration Test</h1>
<p>This HTML should convert to text and PDF through the web interface.</p>
</body></html>"""
    with open("web_test.html", 'w', encoding='utf-8') as f:
        f.write(html_content)
    files['html'] = "web_test.html"
    
    # Create test DOCX
    try:
        import docx
        doc = docx.Document()
        doc.add_heading('Web Integration Test', 0)
        doc.add_paragraph('This DOCX should convert to PDF through the web interface.')
        doc.save("web_test.docx")
        files['docx'] = "web_test.docx"
    except:
        print("⚠️  DOCX creation skipped (python-docx not available)")
    
    return files

def test_api_conversion(input_file, source_format, target_format, test_name):
    """Test a single conversion through the API"""
    print(f"  🔄 Testing {test_name}...")
    
    try:
        with open(input_file, 'rb') as f:
            files = {'file': f}
            data = {
                'source_format': source_format,
                'target_format': target_format
            }
            
            response = requests.post(
                "http://localhost:5000/api/convert",
                files=files,
                data=data,
                timeout=30
            )
        
        if response.status_code == 200:
            result = response.json()
            if result['success']:
                print(f"    ✅ {test_name}: SUCCESS ({result['size']} bytes)")
                
                # Test download
                download_url = f"http://localhost:5000{result['download_url']}"
                download_response = requests.get(download_url, timeout=10)
                
                if download_response.status_code == 200:
                    print(f"    ✅ Download: SUCCESS ({len(download_response.content)} bytes)")
                    return True
                else:
                    print(f"    ❌ Download: FAILED ({download_response.status_code})")
                    return False
            else:
                print(f"    ❌ {test_name}: FAILED - {result.get('error', 'Unknown error')}")
                return False
        else:
            print(f"    ❌ {test_name}: API ERROR ({response.status_code})")
            return False
            
    except Exception as e:
        print(f"    ❌ {test_name}: EXCEPTION - {e}")
        return False

def main():
    """Main web integration test"""
    print("🌐 FileAlchemy Web Integration Test")
    print("=" * 50)
    
    # Check API health
    try:
        response = requests.get("http://localhost:5000/api/health", timeout=5)
        if response.status_code == 200:
            print("✅ API server is healthy")
        else:
            print("❌ API server health check failed")
            print("💡 Start the API server: python api_server.py")
            return False
    except Exception as e:
        print(f"❌ Cannot connect to API server: {e}")
        print("💡 Start the API server: python api_server.py")
        return False
    
    # Create test files
    print("\n📁 Creating test files...")
    test_files = create_test_files()
    
    # Test key conversions
    print("\n🧪 Testing Key Web Conversions:")
    results = {}
    
    # Text conversions
    if 'txt' in test_files:
        results['TXT→HTML'] = test_api_conversion(test_files['txt'], 'TXT', 'HTML', 'TXT to HTML')
        results['TXT→PDF'] = test_api_conversion(test_files['txt'], 'TXT', 'PDF', 'TXT to PDF')
    
    # HTML conversions
    if 'html' in test_files:
        results['HTML→TXT'] = test_api_conversion(test_files['html'], 'HTML', 'TXT', 'HTML to TXT')
        results['HTML→PDF'] = test_api_conversion(test_files['html'], 'HTML', 'PDF', 'HTML to PDF')
    
    # DOCX conversions
    if 'docx' in test_files:
        results['DOCX→PDF'] = test_api_conversion(test_files['docx'], 'DOCX', 'PDF', 'DOCX to PDF')
        results['DOCX→TXT'] = test_api_conversion(test_files['docx'], 'DOCX', 'TXT', 'DOCX to TXT')
    
    # Results summary
    print("\n" + "=" * 50)
    print("📊 Web Integration Test Results:")
    print("-" * 30)
    
    total_tests = len(results)
    successful_tests = sum(1 for success in results.values() if success)
    
    for test_name, success in results.items():
        status = "✅ PASSED" if success else "❌ FAILED"
        print(f"   {test_name}: {status}")
    
    print("-" * 30)
    print(f"   Overall: {successful_tests}/{total_tests} web tests passed")
    
    # Cleanup
    print("\n🧹 Cleaning up test files...")
    for file in test_files.values():
        if os.path.exists(file):
            os.remove(file)
            print(f"   Removed: {file}")
    
    if successful_tests == total_tests:
        print("\n🎉 Web integration tests PASSED!")
        print("🚀 FileAlchemy is ready for web deployment!")
        return True
    else:
        print(f"\n⚠️  {total_tests - successful_tests} web tests failed")
        print("💡 Check the API server and try again")
        return False

if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)